"""
Image to DOCX Generator
Generate DOCX document from test method images with Excel references.

Usage:
    python image_to_docx.py --images folder_path --json test_data.json --output output.docx
    python image_to_docx.py --images folder_path --json test_data.json --output output.docx --append existing.docx
"""

import argparse
import json
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_test_mapping(json_file: str) -> dict:
    """
    Load JSON and create mapping from method name to sheet/TC info.
    
    Returns dict: {method_name: {"sheet": sheet_name, "tc_id": TC01, "display_name": ...}}
    """
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    if not isinstance(data, list):
        data = [data]
    
    mapping = {}
    
    for scenario in data:
        sheet_name = scenario['metadata']['sheet_name']
        
        for tc in scenario['test_cases']:
            method_name = tc.get('note', '')  # method name stored in note
            if method_name:
                mapping[method_name] = {
                    'sheet': sheet_name,
                    'tc_id': tc['id'],
                    'display_name': tc.get('display_name', method_name),
                    'type': tc.get('type', 'N')
                }
    
    return mapping


def parse_image_filename(filename: str) -> tuple:
    """
    Parse image filename to extract nested class and method name.
    
    Format: NestedClass_methodName.png or methodName.png
    Returns: (nested_class, method_name)
    """
    name = Path(filename).stem  # Remove extension
    
    # Check if has nested class prefix (contains underscore)
    if '_' in name:
        parts = name.split('_', 1)
        return parts[0], parts[1]
    
    return None, name


def create_docx_from_images(
    images_dir: str,
    json_file: str,
    output_file: str,
    append_to: str = None,
    title: str = None
) -> str:
    """
    Create DOCX document from test images with Excel references.
    
    Args:
        images_dir: Directory containing test method images
        json_file: JSON file with test case mapping
        output_file: Output DOCX file path
        append_to: Existing DOCX file to append to (optional)
        title: Document/section title (default: derived from json filename)
    
    Returns:
        Output file path
    """
    # Load mapping
    mapping = load_test_mapping(json_file)
    
    # Get all images
    image_files = sorted([
        f for f in os.listdir(images_dir) 
        if f.lower().endswith(('.png', '.jpg', '.jpeg'))
    ])
    
    if not image_files:
        print(f"[WARN] No images found in {images_dir}")
        return None
    
    # Create or open document
    if append_to and os.path.exists(append_to):
        doc = Document(append_to)
        # Add page break before new content
        doc.add_page_break()
    else:
        doc = Document()
    
    # Derive title from json filename if not provided
    if not title:
        title = Path(json_file).stem.replace('_', ' ').title()
    
    # Add title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Group images by nested class
    grouped = {}
    for img_file in image_files:
        nested_class, method_name = parse_image_filename(img_file)
        key = nested_class or '_root'
        if key not in grouped:
            grouped[key] = []
        grouped[key].append((img_file, method_name))
    
    # Process each group
    image_count = 0
    for nested_class, images in grouped.items():
        # Add nested class heading if not root
        if nested_class != '_root':
            doc.add_heading(nested_class, level=2)
        
        for img_file, method_name in images:
            image_count += 1
            img_path = os.path.join(images_dir, img_file)
            
            # Get mapping info
            info = mapping.get(method_name, {})
            sheet = info.get('sheet', 'N/A')
            tc_id = info.get('tc_id', 'N/A')
            
            # Add image
            doc.add_picture(img_path, width=Inches(6.5))
            
            # Add caption with reference
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = caption.add_run(f'Figure: {method_name}')
            run.bold = True
            run.font.size = Pt(10)
            
            caption.add_run('\n')
            
            ref_run = caption.add_run(f'Reference: Sheet "{sheet}", Test Case {tc_id}')
            ref_run.italic = True
            ref_run.font.size = Pt(9)
            
            # Add spacing
            doc.add_paragraph()
    
    # Save document
    doc.save(output_file)
    
    print(f"[OK] Generated DOCX with {image_count} images: {output_file}")
    return output_file


def main():
    parser = argparse.ArgumentParser(
        description='Generate DOCX from test method images with Excel references'
    )
    parser.add_argument('--images', required=True, help='Directory containing test images')
    parser.add_argument('--json', required=True, help='JSON file with test case mapping')
    parser.add_argument('--output', required=True, help='Output DOCX file path')
    parser.add_argument('--append', help='Append to existing DOCX file')
    parser.add_argument('--title', help='Document/section title')
    
    args = parser.parse_args()
    
    if not os.path.isdir(args.images):
        print(f"[ERROR] Images directory not found: {args.images}")
        return 1
    
    if not os.path.exists(args.json):
        print(f"[ERROR] JSON file not found: {args.json}")
        return 1
    
    create_docx_from_images(
        images_dir=args.images,
        json_file=args.json,
        output_file=args.output,
        append_to=args.append,
        title=args.title
    )
    
    return 0


if __name__ == '__main__':
    exit(main())
